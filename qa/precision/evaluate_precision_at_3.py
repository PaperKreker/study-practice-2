import argparse
import csv
import json
import os
import tempfile
import uuid
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode
from urllib.request import Request, urlopen


DEFAULT_QUERIES = Path(__file__).with_name("queries.csv")


def request_json(
    url: str,
    *,
    method: str = "GET",
    token: str | None = None,
    payload: dict | None = None,
) -> dict:
    data = None
    headers = {"Accept": "application/json"}
    if payload is not None:
        data = json.dumps(payload).encode("utf-8")
        headers["Content-Type"] = "application/json"
    if token:
        headers["Authorization"] = f"Bearer {token}"

    request = Request(url, data=data, headers=headers, method=method)
    try:
        with urlopen(request, timeout=30) as response:
            return json.load(response)
    except HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"HTTP {exc.code} from {url}: {body}") from exc
    except URLError as exc:
        raise RuntimeError(f"Cannot connect to {url}: {exc.reason}") from exc


def get_token(base_url: str) -> str:
    token = os.getenv("QA_ACCESS_TOKEN")
    if token:
        return token

    username = os.getenv("QA_USERNAME", "qa-precision")
    password = os.getenv("QA_PASSWORD", "qa-precision-password")
    credentials = {"username": username, "password": password}
    try:
        response = request_json(
            f"{base_url}/api/v1/users/login",
            method="POST",
            payload=credentials,
        )
    except RuntimeError as login_error:
        try:
            response = request_json(
                f"{base_url}/api/v1/users/register",
                method="POST",
                payload=credentials,
            )
        except RuntimeError:
            raise login_error
    try:
        return response["access_token"]
    except KeyError as exc:
        raise RuntimeError("Login response does not contain access_token") from exc


def load_queries(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as source:
        rows = list(csv.DictReader(source))

    if len(rows) != 10:
        raise ValueError(f"Expected exactly 10 reference queries, got {len(rows)}")
    required_fields = ("query", "expected_file", "content")
    if any(not all(row.get(field) for field in required_fields) for row in rows):
        raise ValueError(
            "Each reference row must contain query, expected_file, and content"
        )
    return rows


def upload_document(base_url: str, token: str, path: Path) -> None:
    boundary = f"----qa-precision-{uuid.uuid4().hex}"
    file_bytes = path.read_bytes()
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{path.name}"\r\n'
        "Content-Type: application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document\r\n\r\n"
    ).encode("ascii")
    body += file_bytes + f"\r\n--{boundary}--\r\n".encode("ascii")
    request = Request(
        f"{base_url}/api/v1/documents/upload",
        data=body,
        headers={
            "Accept": "application/json",
            "Authorization": f"Bearer {token}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
        method="POST",
    )
    try:
        with urlopen(request, timeout=60) as response:
            json.load(response)
    except HTTPError as exc:
        response_body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(
            f"Cannot upload {path.name}: HTTP {exc.code}: {response_body}"
        ) from exc
    except URLError as exc:
        raise RuntimeError(f"Cannot upload {path.name}: {exc.reason}") from exc


def prepare_dataset(base_url: str, token: str, rows: list[dict[str, str]]) -> None:
    from docx import Document

    params = urlencode({"limit": 200, "my_docs": "true"})
    response = request_json(
        f"{base_url}/api/v1/documents?{params}",
        token=token,
    )
    existing = {
        item.get("file_name", "").casefold() for item in response.get("items", [])
    }

    with tempfile.TemporaryDirectory(prefix="qa-precision-") as temp_dir:
        directory = Path(temp_dir)
        for row in rows:
            file_name = row["expected_file"]
            if file_name.casefold() in existing:
                print(f"SKIP: {file_name} is already uploaded")
                continue

            document = Document()
            document.add_heading(file_name.removesuffix(".docx"), level=1)
            document.add_paragraph(row["content"])
            path = directory / file_name
            document.save(path)
            upload_document(base_url, token, path)
            print(f"UPLOAD: {file_name}")


def evaluate(base_url: str, token: str, rows: list[dict[str, str]]) -> list[dict]:
    results = []
    for row in rows:
        query = row["query"]
        expected = row["expected_file"]
        params = urlencode({"q": query, "page": 1, "size": 3})
        response = request_json(
            f"{base_url}/api/v1/search?{params}",
            token=token,
        )
        top_files = [
            item.get("file_name", "") for item in response.get("items", [])[:3]
        ]
        hit = expected.casefold() in {name.casefold() for name in top_files}
        results.append(
            {
                "query": query,
                "expected_file": expected,
                "top_3": " | ".join(top_files),
                "hit": hit,
            }
        )
    return results


def write_results(path: Path, results: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as target:
        writer = csv.DictWriter(
            target,
            fieldnames=["query", "expected_file", "top_3", "hit"],
        )
        writer.writeheader()
        writer.writerows(results)


def main() -> int:
    parser = argparse.ArgumentParser(description="Evaluate search Precision@3")
    parser.add_argument("--base-url", default="http://localhost:8000")
    parser.add_argument("--queries", type=Path, default=DEFAULT_QUERIES)
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).with_name("report.csv"),
    )
    parser.add_argument(
        "--prepare-dataset",
        action="store_true",
        help="Create and upload the ten deterministic reference documents",
    )
    args = parser.parse_args()

    rows = load_queries(args.queries)
    token = get_token(args.base_url.rstrip("/"))
    if args.prepare_dataset:
        prepare_dataset(args.base_url.rstrip("/"), token, rows)
    results = evaluate(args.base_url.rstrip("/"), token, rows)
    write_results(args.output, results)

    hits = sum(result["hit"] for result in results)
    precision = hits / len(results)
    for result in results:
        marker = "PASS" if result["hit"] else "FAIL"
        print(f"{marker}: {result['query']} -> {result['top_3']}")
    print(f"Precision@3: {hits}/{len(results)} = {precision:.2%}")
    print(f"Report: {args.output}")
    return 0 if hits == len(results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
